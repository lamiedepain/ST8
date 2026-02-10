#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ST8 Planning - Application Flask
Application web pour la gestion du planning et des équipes ST8
"""

from flask import Flask, render_template, jsonify, request, send_file
from datetime import datetime, timedelta
from openpyxl import load_workbook as openpyxl_load_workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from pathlib import Path
import os
from dotenv import load_dotenv
import io

# Charger variables d'environnement
load_dotenv()

# Import modules locaux
from config import *
from teams_structure import AGENTS, get_agent_equipe, get_agent_equipe_color
from utils import (
    load_workbook, save_workbook, create_backup, cell_to_str, cell_to_date_str, normalize_status,
    find_header_row, list_excel_files, resolve_excel_path, list_backups,
    restore_backup, ExcelError
)


# ==============================================================================
# INITIALISATION FLASK
# ==============================================================================

app = Flask(__name__)
app.config['SECRET_KEY'] = SECRET_KEY
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # Max 50MB upload

# ==============================================================================
# VÉRIFICATION AU DÉMARRAGE
# ==============================================================================

def startup_sync():
    """Vérifie le fichier au démarrage"""
    print("\n" + "="*60)
    print("ST8 PLANNING - Démarrage")
    print("="*60)
    
    # Vérifier fichier local
    if LOCAL_EXCEL_PATH.exists():
        print(f"✅ Fichier Excel: {EXCEL_FILE}")
    else:
        print(f"⚠️  Fichier Excel introuvable: {EXCEL_FILE}")
        print("⚠️  Mode dégradé - Uploadez un fichier Excel via l'interface")
    
    print("\n" + "="*60)

# ==============================================================================
# ROUTES - PAGES HTML
# ==============================================================================

@app.route('/')
def index():
    """Page d'accueil"""
    return render_template('index.html', 
                         excel_file=EXCEL_FILE)

@app.route('/planning')
def planning_page():
    """Page de planification"""
    return render_template('planning.html')

@app.route('/generator')
def generator_page():
    """Page générateur de dossiers"""
    # Charger les articles du magasin
    try:
        import json
        magasin_path = Path(__file__).parent / 'articles_magasin.json'
        if magasin_path.exists():
            with open(magasin_path, 'r', encoding='utf-8') as f:
                articles = json.load(f)
        else:
            articles = []
    except:
        articles = []
    
    return render_template('generator.html', articles_json=json.dumps(articles))

@app.route('/agents')
def agents_page():
    """Page de gestion des agents"""
    return render_template('agents.html')

# ==============================================================================
# ROUTES - EASYDICT
# ==============================================================================

@app.route('/easydict/<path:filename>')
def easydict_files(filename):
    """Servir les fichiers du dossier easydict"""
    easydict_dir = str(Path(__file__).parent / 'easydict')
    return send_file(Path(easydict_dir) / filename)

@app.route('/api/easydict-template')
def get_easydict_template():
    """Servir le template XLSM pour EasyDict"""
    template_path = Path(__file__).parent / 'easydict' / 'T-DICT_Template.xlsm'
    if template_path.exists():
        return send_file(template_path, mimetype='application/vnd.ms-excel.sheet.macroEnabled.12')
    else:
        return jsonify({'error': 'Template non trouvé'}), 404

@app.route('/api/magasin-articles')
def get_magasin_articles():
    """Récupérer la liste des articles du magasin"""
    try:
        magasin_path = Path(__file__).parent / 'magasin_st8' / 'magasin.xlsx'
        if not magasin_path.exists():
            return jsonify({'success': False, 'error': 'Fichier magasin non trouvé'}), 404
        
        wb = openpyxl_load_workbook(magasin_path, read_only=True, data_only=True)
        ws = wb.active
        
        articles = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1]:  # Article et Description
                articles.append({
                    'code': row[0],
                    'description': row[1],
                    'stock': row[6] if len(row) > 6 else 0
                })
        
        wb.close()
        return jsonify({'success': True, 'articles': articles, 'count': len(articles)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/rues-st8')
def get_rues_st8():
    """Récupérer la liste des rues de ST8"""
    try:
        rues_path = Path(__file__).parent / 'rues_st8.json'
        if not rues_path.exists():
            return jsonify({'success': False, 'error': 'Fichier rues non trouvé'}), 404
        
        with open(rues_path, 'r', encoding='utf-8') as f:
            rues = json.load(f)
        
        return jsonify({'success': True, 'rues': rues, 'count': len(rues)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/outillage')
def get_outillage():
    """Récupérer la liste de l'outillage DPM"""
    try:
        outil_path = Path(__file__).parent / 'magasin_st8' / 'outil_dpm.xlsx'
        if not outil_path.exists():
            return jsonify({'success': False, 'error': 'Fichier outillage non trouvé'}), 404
        
        wb = openpyxl_load_workbook(outil_path, read_only=True, data_only=True)
        ws = wb.active
        
        outils = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1]:  # Article et Description
                outils.append({
                    'code': row[0],
                    'description': row[1],
                    'categorie': row[2] if len(row) > 2 else '',
                    'stock': row[6] if len(row) > 6 else 0
                })
        
        wb.close()
        return jsonify({'success': True, 'outils': outils, 'count': len(outils)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/vehicules')
def get_vehicules():
    """Récupérer la liste des véhicules"""
    try:
        vehicules_path = Path(__file__).parent / 'magasin_st8' / 'vehicules.xlsx'
        if not vehicules_path.exists():
            return jsonify({'success': False, 'error': 'Fichier véhicules non trouvé'}), 404
        
        wb = openpyxl_load_workbook(vehicules_path, read_only=True, data_only=True)
        ws = wb.active
        
        vehicules = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] or row[1] or row[2]:  # Code mage, Véhicule ou Immatriculation
                vehicules.append({
                    'code_mage': str(row[0]) if row[0] else '',
                    'numero': row[1] if row[1] else '',
                    'immatriculation': row[2] if len(row) > 2 and row[2] else ''
                })
        
        wb.close()
        return jsonify({'success': True, 'vehicules': vehicules, 'count': len(vehicules)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/agents-disponibles', methods=['POST'])
def get_agents_disponibles():
    """Récupérer les agents voirie disponibles pour une période"""
    try:
        data = request.json
        date_debut = data.get('dateDebut')
        date_fin = data.get('dateFin')
        
        if not date_debut or not date_fin:
            return jsonify({'success': False, 'error': 'Dates manquantes'}), 400
        
        # Convertir les dates
        debut = datetime.strptime(date_debut, '%Y-%m-%d')
        fin = datetime.strptime(date_fin, '%Y-%m-%d')
        
        # Charger le fichier Excel
        wb = load_workbook()
        
        # Filtrer uniquement les agents de voirie (pas les responsables)
        agents_voirie = [a for a in AGENTS if get_agent_equipe(a['nom']) == 'Agent.es de voirie']
        agents_disponibles = []
        
        # Pour chaque agent voirie
        for agent in agents_voirie:
            est_disponible = True
            agent_trouve_dans_excel = False
            
            # Vérifier les mois entre debut et fin
            current_date = debut
            while current_date <= fin and est_disponible:
                # Ignorer les weekends (5=samedi, 6=dimanche)
                if current_date.weekday() >= 5:
                    current_date += timedelta(days=1)
                    continue
                
                sheet_name = get_month_sheet_name(current_date.year, current_date.month)
                
                # Vérifier si la feuille existe
                if sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    header_row = find_header_row(sheet, 'Matricule')
                    
                    if header_row:
                        agent_trouve_ce_jour = False
                        # Trouver la ligne de l'agent
                        for row_idx in range(header_row + 1, sheet.max_row + 1):
                            nom = cell_to_str(sheet.cell(row_idx, 2).value)
                            prenom = cell_to_str(sheet.cell(row_idx, 3).value)
                            
                            if nom and prenom:
                                if nom.strip().upper() == agent['nom'].upper() and prenom.strip().upper() == agent['prenom'].upper():
                                    agent_trouve_ce_jour = True
                                    agent_trouve_dans_excel = True
                                    
                                    # Vérifier le jour spécifique
                                    jour = current_date.day
                                    col_idx = PLANNING_DAYS_START_COL + jour - 1
                                    statut = normalize_status(sheet.cell(row_idx, col_idx).value)
                                    
                                    # Si statut indique une absence
                                    if statut and statut.upper() in ABSENT_STATUSES:
                                        est_disponible = False
                                    break
                        
                        # Si l'agent n'est pas trouvé dans le planning, on ne peut pas garantir sa disponibilité
                        if not agent_trouve_ce_jour:
                            est_disponible = False
                
                # Passer au jour suivant
                current_date += timedelta(days=1)
            
            # Ajouter seulement si l'agent est disponible ET trouvé dans l'Excel
            if est_disponible and agent_trouve_dans_excel:
                agents_disponibles.append(agent)
        
        wb.close()
        
        return jsonify({'success': True, 'agents': agents_disponibles, 'count': len(agents_disponibles)})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/generate-reference', methods=['POST'])
def generate_reference():
    """Générer une référence automatique pour un chantier"""
    try:
        data = request.json
        quartier = data.get('quartier', 'Q1')
        date_debut = data.get('dateDebut')
        
        # Utiliser la date de début si fournie, sinon la date actuelle
        if date_debut:
            ref_date = datetime.strptime(date_debut, '%Y-%m-%d')
        else:
            ref_date = datetime.now()
        
        annee = ref_date.year
        mois = f"{ref_date.month:02d}"
        
        # Charger les références existantes
        refs_path = Path(__file__).parent / 'chantiers_references.json'
        if refs_path.exists():
            with open(refs_path, 'r', encoding='utf-8') as f:
                references = json.load(f)
        else:
            references = {}
        
        # Clé pour ce quartier et ce mois
        key = f"{quartier}-{annee}-{mois}"
        
        # Trouver le prochain numéro
        if key in references:
            dernier_numero = references[key]
            prochain_numero = dernier_numero + 1
        else:
            prochain_numero = 1
        
        # Mettre à jour les références
        references[key] = prochain_numero
        
        # Sauvegarder
        with open(refs_path, 'w', encoding='utf-8') as f:
            json.dump(references, f, indent=2, ensure_ascii=False)
        
        # Générer la référence
        reference = f"{quartier}-{annee}-{mois}-{prochain_numero:03d}"
        
        return jsonify({
            'success': True,
            'reference': reference
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/generate-fiche-docx', methods=['POST'])
def generate_fiche_docx():
    """Générer un document Word formaté pour la fiche travaux"""
    try:
        data = request.json
        
        # Créer un nouveau document
        doc = Document()
        
        # Définir les marges (optimisées pour A4)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # === LOGO ET EN-TÊTE ===
        # Essayer d'ajouter le logo
        logo_path = Path(__file__).parent / 'static' / 'images' / 'logo_bordeaux_metropole.png'
        if logo_path.exists():
            header_para = doc.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_para.add_run()
            run.add_picture(str(logo_path), width=Inches(1.3))
        else:
            # Fallback si pas de logo
            header = doc.add_heading('BORDEAUX MÉTROPOLE', level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            header.runs[0].font.size = Pt(20)
            header.runs[0].font.bold = True
        
        subtitle = doc.add_paragraph('Régie Voirie Espaces Verts | 60 rue New-York')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(8)
        subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
        
        # === TITRE FICHE ===
        titre = doc.add_heading('FICHE TRAVAUX', level=1)
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titre.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        titre.runs[0].font.size = Pt(13)
        titre.runs[0].font.bold = True
        titre.space_after = Pt(3)
        
        # Nom du chantier
        nom_chantier = doc.add_paragraph()
        nom_chantier.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nom_run = nom_chantier.add_run(data.get('nom', ''))
        nom_run.font.size = Pt(11)
        nom_run.font.bold = True
        nom_run.font.color.rgb = RGBColor(51, 51, 51)
        nom_chantier.space_after = Pt(6)
        
        # === LOCALISATION ===
        loc_heading = doc.add_heading('📍 LOCALISATION', level=2)
        loc_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        loc_heading.runs[0].font.size = Pt(10)
        loc_heading.space_before = Pt(3)
        loc_heading.space_after = Pt(3)
        
        table_loc = doc.add_table(rows=4, cols=2)
        table_loc.style = 'Medium Grid 1 Accent 1'
        table_loc.autofit = True
        
        cells_loc = [
            ('Quartier', data.get('quartier', '')),
            ('Référence', data.get('reference', '')),
            ('Adresse', f"{data.get('numero', '')} {data.get('adresse', '')}".strip()),
            ('Responsable', data.get('responsable', ''))
        ]
        
        for i, (label, value) in enumerate(cells_loc):
            cell_label = table_loc.rows[i].cells[0]
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].font.bold = True
            cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
            cell_label.paragraphs[0].runs[0].font.size = Pt(9)
            table_loc.rows[i].cells[1].text = value
            table_loc.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        
        # === TYPE DE DEMANDE ===
        if any([data.get('numeroGDU'), data.get('numeroMages'), data.get('dateDemande'), data.get('nomDemandeur')]):
            demande_heading = doc.add_heading('📋 TYPE DE DEMANDE', level=2)
            demande_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            demande_heading.runs[0].font.size = Pt(10)
            demande_heading.space_before = Pt(3)
            demande_heading.space_after = Pt(3)
            
            table_demande = doc.add_table(rows=4, cols=2)
            table_demande.style = 'Medium Grid 1 Accent 1'
            
            cells_demande = [
                ('N° GDU', data.get('numeroGDU', '')),
                ('N° Mages', data.get('numeroMages', '')),
                ('Date de demande', data.get('dateDemande', '')),
                ('Demandeur', data.get('nomDemandeur', ''))
            ]
            
            for i, (label, value) in enumerate(cells_demande):
                cell_label = table_demande.rows[i].cells[0]
                cell_label.text = label
                cell_label.paragraphs[0].runs[0].font.bold = True
                cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
                cell_label.paragraphs[0].runs[0].font.size = Pt(9)
                table_demande.rows[i].cells[1].text = value
                table_demande.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        
        # === DATES ET DURÉE ===
        dates_heading = doc.add_heading('📅 DATES ET DURÉE', level=2)
        dates_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        dates_heading.runs[0].font.size = Pt(10)
        dates_heading.space_before = Pt(3)
        dates_heading.space_after = Pt(3)
        
        table_dates = doc.add_table(rows=3, cols=2)
        table_dates.style = 'Medium Grid 1 Accent 1'
        
        cells_dates = [
            ('Date de début', data.get('dateDebut', '')),
            ('Date de fin prévisionnelle', data.get('dateFin', '')),
            ('Nombre de jours envisagés', data.get('nombreJours', ''))
        ]
        
        for i, (label, value) in enumerate(cells_dates):
            cell_label = table_dates.rows[i].cells[0]
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].font.bold = True
            cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
            cell_label.paragraphs[0].runs[0].font.size = Pt(9)
            table_dates.rows[i].cells[1].text = str(value) if value else ''
            table_dates.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        
        # === PRÉPARATION ===
        prep_heading = doc.add_heading('🔍 PRÉPARATION', level=2)
        prep_heading.runs[0].font.color.rgb = RGBColor(0, 153, 76)
        prep_heading.runs[0].font.size = Pt(10)
        prep_heading.space_before = Pt(3)
        prep_heading.space_after = Pt(3)
        
        table_prep = doc.add_table(rows=5, cols=2)
        table_prep.style = 'Medium List 1 Accent 3'
        
        cells_prep = [
            ('Visite sur site', data.get('visiteSite', '')),
            ('Date de visite', data.get('dateVisite', '')),
            ("Nom de l'AM", data.get('nomAM', '')),
            ('Travaux envisagés', data.get('travauxEnvisages', '')),
            ('Priorisation', data.get('priorisationLabel', ''))
        ]
        
        for i, (label, value) in enumerate(cells_prep):
            cell_label = table_prep.rows[i].cells[0]
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].font.bold = True
            cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 51)
            cell_label.paragraphs[0].runs[0].font.size = Pt(9)
            table_prep.rows[i].cells[1].text = value
            table_prep.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        
        # === PRÉALABLES ===
        prealables_heading = doc.add_heading('⚠️ PRÉALABLES', level=2)
        prealables_heading.runs[0].font.color.rgb = RGBColor(255, 140, 0)
        prealables_heading.runs[0].font.size = Pt(10)
        prealables_heading.space_before = Pt(3)
        prealables_heading.space_after = Pt(3)
        
        table_prealables = doc.add_table(rows=8, cols=2)
        table_prealables.style = 'Medium List 1 Accent 6'
        
        cells_prealables = [
            ('Demande Amiante', data.get('demandeAmiante', '')),
            ('Présence Amiante', data.get('presenceAmiante', '')),
            ('N° DT/DICT', data.get('numeroDTDICT', '')),
            ('Date DICT', data.get('dateDICT', '')),
            ('Arrêté demandé le', data.get('dateDemandeArrete', '')),
            ("Date de l'arrêté", data.get('dateArrete', '')),
            ('Proximité Tram', data.get('proximiteTram', '')),
            ('Protocole signé', data.get('protocoleSign', ''))
        ]
        
        for i, (label, value) in enumerate(cells_prealables):
            cell_label = table_prealables.rows[i].cells[0]
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].font.bold = True
            cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(153, 76, 0)
            cell_label.paragraphs[0].runs[0].font.size = Pt(9)
            table_prealables.rows[i].cells[1].text = value
            table_prealables.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        
        # === DESCRIPTION DU CHANTIER ===
        if data.get('description'):
            desc_heading = doc.add_heading('📝 DESCRIPTION', level=2)
            desc_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            desc_heading.runs[0].font.size = Pt(10)
            desc_heading.space_before = Pt(3)
            desc_heading.space_after = Pt(3)
            
            desc_para = doc.add_paragraph(data.get('description', ''))
            desc_para.runs[0].font.size = Pt(9)
        
        # === RESSOURCES ===
        ressources_heading = doc.add_heading('👥 RESSOURCES', level=2)
        ressources_heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        ressources_heading.runs[0].font.size = Pt(10)
        ressources_heading.space_before = Pt(3)
        ressources_heading.space_after = Pt(3)
        
        table_ressources = doc.add_table(rows=3, cols=2)
        table_ressources.style = 'Medium Grid 1 Accent 1'
        
        cells_ressources = [
            ("Nombre d'agents", data.get('nombreAgents', '')),
            ('Matériaux nécessaires', data.get('materiauxNecessaires', '')),
            ('Permis/CACES', data.get('permisCacesNecessaires', ''))
        ]
        
        for i, (label, value) in enumerate(cells_ressources):
            cell_label = table_ressources.rows[i].cells[0]
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].font.bold = True
            cell_label.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
            cell_label.paragraphs[0].runs[0].font.size = Pt(9)
            table_ressources.rows[i].cells[1].text = value
            table_ressources.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        
        # === AGENTS, VÉHICULES, MATÉRIAUX ===
        if data.get('agents'):
            agents_heading = doc.add_heading('Agents affectés', level=3)
            agents_heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            agents_heading.runs[0].font.size = Pt(9)
            agents_heading.space_before = Pt(2)
            agents_heading.space_after = Pt(2)
            for agent in data.get('agents', []):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(agent).font.size = Pt(8)
                p.space_after = Pt(1)
        
        if data.get('vehicules'):
            veh_heading = doc.add_heading('Véhicules', level=3)
            veh_heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            veh_heading.runs[0].font.size = Pt(9)
            veh_heading.space_before = Pt(2)
            veh_heading.space_after = Pt(2)
            for veh in data.get('vehicules', []):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(veh).font.size = Pt(8)
                p.space_after = Pt(1)
        
        if data.get('articles'):
            art_heading = doc.add_heading('Matériaux', level=3)
            art_heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            art_heading.runs[0].font.size = Pt(9)
            art_heading.space_before = Pt(2)
            art_heading.space_after = Pt(2)
            for art in data.get('articles', []):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(art).font.size = Pt(8)
                p.space_after = Pt(1)
        
        if data.get('outils'):
            outil_heading = doc.add_heading('Outillage', level=3)
            outil_heading.runs[0].font.color.rgb = RGBColor(51, 102, 153)
            outil_heading.runs[0].font.size = Pt(9)
            outil_heading.space_before = Pt(2)
            outil_heading.space_after = Pt(2)
            for outil in data.get('outils', []):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(outil).font.size = Pt(8)
                p.space_after = Pt(1)
        
        # === PIED DE PAGE ===
        footer = doc.add_paragraph(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(7)
        footer.runs[0].font.italic = True
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        footer.space_before = Pt(6)
        
        # Sauvegarder en mémoire
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Nom du fichier
        nom_fichier = data.get('nom', 'Fiche_travaux').replace(' ', '_').replace('/', '_')
        filename = f"Fiche_Travaux_{nom_fichier}.docx"
        
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# ==============================================================================
# ROUTES - FICHIERS
# ==============================================================================

@app.route('/api/files')
def get_files():
    """
    Liste les fichiers Excel disponibles
    
    Returns:
        JSON: {files: [liste fichiers], backups: [liste backups]}
    """
    try:
        files = list_excel_files()
        backups = list_backups()
        
        return jsonify({
            'success': True,
            'files': files,
            'backups': backups,
            'current_file': str(LOCAL_EXCEL_PATH),
            'sharepoint_enabled': AUTO_SYNC_ENABLED
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/download-excel')
def download_excel():
    """
    Télécharge le fichier Excel principal
    
    Returns:
        File: Fichier Excel
    """
    try:
        if not LOCAL_EXCEL_PATH.exists():
            return jsonify({
                'success': False,
                'error': 'Fichier Excel introuvable'
            }), 404
        
        return send_file(
            LOCAL_EXCEL_PATH,
            as_attachment=True,
            download_name=EXCEL_FILE
        )
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/reload-excel', methods=['POST'])
def reload_excel():
    """
    Recharge les métadonnées du fichier Excel
    
    Returns:
        JSON: {success: bool, sheets: [liste feuilles], agents: int}
    """
    try:
        wb = load_workbook()
        
        # Compter agents
        agents_count = 0
        if 'config' in wb.sheetnames:
            config_sheet = wb['config']
            header_row = find_header_row(config_sheet, 'Matricule')
            if header_row:
                for row_idx in range(header_row + 1, config_sheet.max_row + 1):
                    if config_sheet.cell(row_idx, 1).value:
                        agents_count += 1
        
        return jsonify({
            'success': True,
            'sheets': wb.sheetnames,
            'agents_count': agents_count,
            'file': EXCEL_FILE
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/load-planning', methods=['POST'])
def load_planning():
    """
    Upload un nouveau fichier de planning
    
    Returns:
        JSON: {success: bool, message: str}
    """
    try:
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'Aucun fichier fourni'
            }), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'Nom de fichier vide'
            }), 400
        
        # Vérifier extension
        if not file.filename.endswith(('.xlsx', '.xlsm')):
            return jsonify({
                'success': False,
                'error': 'Fichier doit être .xlsx ou .xlsm'
            }), 400
        
        # Créer backup de l'ancien fichier
        if LOCAL_EXCEL_PATH.exists():
            create_backup()
        
        # Sauvegarder nouveau fichier
        file.save(LOCAL_EXCEL_PATH)
        
        return jsonify({
            'success': True,
            'message': 'Fichier chargé avec succès'
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==============================================================================
# ROUTES - BACKUPS
# ==============================================================================

@app.route('/api/backups')
def get_backups():
    """
    Liste tous les backups disponibles
    
    Returns:
        JSON: {success: bool, backups: [liste]}
    """
    try:
        backups = list_backups()
        return jsonify({
            'success': True,
            'backups': backups
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/restore-backup', methods=['POST'])
def api_restore_backup():
    """
    Restaure un backup spécifique
    
    Body: {backup_name: str}
    
    Returns:
        JSON: {success: bool, message: str}
    """
    try:
        data = request.get_json()
        backup_name = data.get('backup_name')
        
        if not backup_name:
            return jsonify({
                'success': False,
                'error': 'Nom de backup manquant'
            }), 400
        
        restore_backup(backup_name)
        
        return jsonify({
            'success': True,
            'message': f'Backup {backup_name} restauré'
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==============================================================================
# ROUTES - AGENTS
# ==============================================================================

def parse_date_echeance(date_str):
    """
    Parse une date d'échéance dans plusieurs formats possibles
    
    Formats supportés:
    - JJ/MM/AAAA (15/05/2026)
    - mai-26, jan-27
    - mai 26, jan 27
    
    Returns:
        datetime ou None
    """
    if not date_str or date_str == '-':
        return None
    
    date_str = str(date_str).strip()
    
    try:
        # Format JJ/MM/AAAA
        if '/' in date_str:
            return datetime.strptime(date_str, '%d/%m/%Y')
        
        # Format mai-26 ou mai 26
        parts = date_str.lower().replace(' ', '-').split('-')
        if len(parts) == 2:
            mois_map = {
                'jan': 1, 'janvier': 1,
                'fev': 2, 'février': 2, 'fevrier': 2,
                'mar': 3, 'mars': 3,
                'avr': 4, 'avril': 4,
                'mai': 5,
                'jui': 6, 'juin': 6,
                'jul': 7, 'juillet': 7,
                'aou': 8, 'août': 8, 'aout': 8,
                'sep': 9, 'septembre': 9,
                'oct': 10, 'octobre': 10,
                'nov': 11, 'novembre': 11,
                'dec': 12, 'décembre': 12, 'decembre': 12
            }
            
            mois_str = parts[0][:3]
            annee_str = parts[1]
            
            # Année sur 2 chiffres (26 -> 2026)
            if len(annee_str) == 2:
                annee = 2000 + int(annee_str)
            else:
                annee = int(annee_str)
            
            if mois_str in mois_map:
                return datetime(annee, mois_map[mois_str], 1)
    except:
        pass
    
    return None

@app.route('/api/agents')
def get_agents():
    """
    Récupère la liste de tous les agents depuis la feuille 'config'
    
    Returns:
        JSON: {success: bool, agents: [liste agents], alerts: {}}
    """
    try:
        wb = load_workbook()
        
        if 'config' not in wb.sheetnames:
            return jsonify({
                'success': False,
                'error': "Feuille 'config' introuvable"
            }), 404
        
        config_sheet = wb['config']
        header_row = find_header_row(config_sheet, 'Matricule')
        
        if not header_row:
            return jsonify({
                'success': False,
                'error': "En-tête 'Matricule' introuvable"
            }), 404
        
        agents = []
        alerts_depassees = []  # Alertes dépassées
        alerts_3_mois = []     # Alertes dans les 3 mois
        alerts_annee = []      # Alertes dans l'année
        
        from datetime import datetime, timedelta
        today = datetime.now()
        date_3_mois = today + timedelta(days=90)
        date_annee = today + timedelta(days=365)
        
        for row_idx in range(header_row + 1, config_sheet.max_row + 1):
            matricule = cell_to_str(config_sheet.cell(row_idx, AGENT_COL_MATRICULE).value)
            
            if matricule:
                nom = cell_to_str(config_sheet.cell(row_idx, AGENT_COL_NOM).value) or ''
                prenom = cell_to_str(config_sheet.cell(row_idx, AGENT_COL_PRENOM).value) or ''
                
                # Lire toutes les colonnes
                agent_equipe = get_agent_equipe(nom)
                agent_equipe_color = get_agent_equipe_color(nom)
                
                agent = {
                    'index': row_idx - header_row,
                    'matricule': matricule,
                    'nom': nom,
                    'prenom': prenom,
                    'equipe': agent_equipe,
                    'equipe_color': agent_equipe_color,
                    'anniversaire': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_ANNIVERSAIRE).value) or '',
                    'dernieres_visites': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_DERNIERES_VISITES).value) or '',
                    'prochaines_visites': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_PROCHAINES_VISITES).value) or '',
                    'r482_a_1': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_A_1).value) or '',
                    'r482_b1_2': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_B1_2).value) or '',
                    'r482_c2_3': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_C2_3).value) or '',
                    'r482_c1_4': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_C1_4).value) or '',
                    'r482_5': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_5).value) or '',
                    'r482_d_7': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_D_7).value) or '',
                    'r482_e_8': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_E_8).value) or '',
                    'r482_f_9': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_F_9).value) or '',
                    'r482_g_10': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CACES_R482_G_10).value) or '',
                    'tondeuse': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_TONDEUSE).value) or '',
                    'grue_r490': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_GRUE_R490).value) or '',
                    'b_nacelle': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_B_NACELLE).value) or '',
                    'chariot_r489_3': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_CHARIOT_R489_3).value) or '',
                    'tronco': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_TRONCO).value) or '',
                    'permis_be': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_PERMIS_BE).value) or '',
                    'permis_pl': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_PERMIS_PL).value) or '',
                    'permis_ce': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_PERMIS_CE).value) or '',
                    'fimo': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_FIMO).value) or '',
                    'aipr': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_AIPR).value) or '',
                    'premiers_secours': cell_to_date_str(config_sheet.cell(row_idx, AGENT_COL_PREMIERS_SECOURS).value) or ''
                }
                
                # Vérifier alertes pour visite médicale
                prochaine_visite = agent['prochaines_visites']
                if prochaine_visite:
                    date_visite = parse_date_echeance(prochaine_visite)
                    if date_visite:
                        if date_visite < today:
                            alerts_depassees.append({
                                'agent': f"{nom} {prenom}",
                                'type': 'Visite médicale',
                                'echeance': prochaine_visite
                            })
                        elif today <= date_visite <= date_3_mois:
                            alerts_3_mois.append({
                                'agent': f"{nom} {prenom}",
                                'type': 'Visite médicale',
                                'echeance': prochaine_visite
                            })
                        elif date_3_mois < date_visite <= date_annee:
                            alerts_annee.append({
                                'agent': f"{nom} {prenom}",
                                'type': 'Visite médicale',
                                'echeance': prochaine_visite
                            })
                
                # Vérifier alertes pour les certifications
                cert_fields = {
                    'R.482 A (1)': agent['r482_a_1'],
                    'R.482 B1 (2)': agent['r482_b1_2'],
                    'R.482 C2 (3)': agent['r482_c2_3'],
                    'R.482 C1 (4)': agent['r482_c1_4'],
                    'R.482 5': agent['r482_5'],
                    'R.482 D (7)': agent['r482_d_7'],
                    'R.482 E (8)': agent['r482_e_8'],
                    'R.482 F (9)': agent['r482_f_9'],
                    'R.482 G (10)': agent['r482_g_10'],
                    'Tondeuse': agent['tondeuse'],
                    'Grue R.490': agent['grue_r490'],
                    'B (Nacelle)': agent['b_nacelle'],
                    'Chariot R.489 / 3': agent['chariot_r489_3'],
                    'Tronço': agent['tronco'],
                    'FIMO': agent['fimo'],
                    'AIPR': agent['aipr'],
                    'Premiers secours': agent['premiers_secours']
                }
                
                for cert_name, cert_value in cert_fields.items():
                    if cert_value and cert_value != '-':
                        date_cert = parse_date_echeance(cert_value)
                        if date_cert:
                            if date_cert < today:
                                alerts_depassees.append({
                                    'agent': f"{nom} {prenom}",
                                    'type': cert_name,
                                    'echeance': cert_value
                                })
                            elif today <= date_cert <= date_3_mois:
                                alerts_3_mois.append({
                                    'agent': f"{nom} {prenom}",
                                    'type': cert_name,
                                    'echeance': cert_value
                                })
                            elif date_3_mois < date_cert <= date_annee:
                                alerts_annee.append({
                                    'agent': f"{nom} {prenom}",
                                    'type': cert_name,
                                    'echeance': cert_value
                                })
                
                agents.append(agent)
        
        return jsonify({
            'success': True,
            'agents': agents,
            'count': len(agents),
            'alerts': {
                'depassees': alerts_depassees,
                '3_mois': alerts_3_mois,
                'annee': alerts_annee
            }
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/agents', methods=['POST'])
def add_agent():
    """
    Ajoute un nouvel agent dans la feuille 'config'
    
    Body: {matricule: str, nom: str, prenom: str, anniversaire: str}
    
    Returns:
        JSON: {success: bool, message: str}
    """
    try:
        data = request.get_json()
        
        # Validation
        matricule = data.get('matricule', '').strip()
        nom = data.get('nom', '').strip()
        prenom = data.get('prenom', '').strip()
        
        if not matricule or not nom:
            return jsonify({
                'success': False,
                'error': 'Matricule et nom obligatoires'
            }), 400
        
        # Créer backup avant modification
        create_backup()
        
        # Charger workbook
        wb = load_workbook(data_only=False)
        config_sheet = wb['config']
        header_row = find_header_row(config_sheet, 'Matricule')
        
        # Trouver première ligne vide
        next_row = header_row + 1
        while config_sheet.cell(next_row, AGENT_COL_MATRICULE).value:
            next_row += 1
        
        # Ajouter agent
        config_sheet.cell(next_row, AGENT_COL_MATRICULE, matricule)
        config_sheet.cell(next_row, AGENT_COL_NOM, nom)
        config_sheet.cell(next_row, AGENT_COL_PRENOM, prenom)
        config_sheet.cell(next_row, AGENT_COL_ANNIVERSAIRE, data.get('anniversaire', ''))
        
        # Sauvegarder
        save_workbook(wb)
        
        # Upload vers SharePoint si activé
        if AUTO_SYNC_ENABLED:
            try:
                upload_to_sharepoint()
            except SharePointError:
                pass  # Continuer même si upload échoue
        
        return jsonify({
            'success': True,
            'message': f'Agent {nom} {prenom} ajouté'
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/agents/<int:index>', methods=['PUT'])
def update_agent(index):
    """
    Met à jour un agent existant
    
    Args:
        index: Index de l'agent (1-based)
    
    Body: {matricule: str, nom: str, prenom: str, anniversaire: str}
    
    Returns:
        JSON: {success: bool, message: str}
    """
    try:
        data = request.get_json()
        
        # Créer backup
        create_backup()
        
        # Charger workbook
        wb = load_workbook(data_only=False)
        config_sheet = wb['config']
        header_row = find_header_row(config_sheet, 'Matricule')
        
        row_idx = header_row + index
        
        # Mettre à jour
        if 'matricule' in data:
            config_sheet.cell(row_idx, AGENT_COL_MATRICULE, data['matricule'])
        if 'nom' in data:
            config_sheet.cell(row_idx, AGENT_COL_NOM, data['nom'])
        if 'prenom' in data:
            config_sheet.cell(row_idx, AGENT_COL_PRENOM, data['prenom'])
        if 'anniversaire' in data:
            config_sheet.cell(row_idx, AGENT_COL_ANNIVERSAIRE, data['anniversaire'])
        
        # Sauvegarder
        save_workbook(wb)
        
        # Upload vers SharePoint
        if AUTO_SYNC_ENABLED:
            try:
                upload_to_sharepoint()
            except SharePointError:
                pass
        
        return jsonify({
            'success': True,
            'message': 'Agent mis à jour'
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/agents/<int:index>', methods=['DELETE'])
def delete_agent(index):
    """
    Supprime un agent
    
    Args:
        index: Index de l'agent (1-based)
    
    Returns:
        JSON: {success: bool, message: str}
    """
    try:
        # Créer backup
        create_backup()
        
        # Charger workbook
        wb = load_workbook(data_only=False)
        config_sheet = wb['config']
        header_row = find_header_row(config_sheet, 'Matricule')
        
        row_idx = header_row + index
        
        # Supprimer ligne
        config_sheet.delete_rows(row_idx, 1)
        
        # Sauvegarder
        save_workbook(wb)
        
        # Upload vers SharePoint
        if AUTO_SYNC_ENABLED:
            try:
                upload_to_sharepoint()
            except SharePointError:
                pass
        
        return jsonify({
            'success': True,
            'message': 'Agent supprimé'
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==============================================================================
# ROUTES - PLANIFICATION
# ==============================================================================

@app.route('/api/months')
def get_months():
    """
    Liste les feuilles mensuelles disponibles
    
    Returns:
        JSON: {success: bool, months: [liste mois]}
    """
    try:
        wb = load_workbook()
        
        # Filtrer feuilles mensuelles (format: "Janvier 2026", etc.)
        months = []
        for sheet_name in wb.sheetnames:
            if any(month_name in sheet_name for month_name in MONTH_NAMES_FR):
                months.append(sheet_name)
        
        return jsonify({
            'success': True,
            'months': months
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/planning-data/<int:year>/<int:month>')
def get_planning_data(year, month):
    """
    Lit les données de planning pour un mois donné
    
    Args:
        year: Année (ex: 2026)
        month: Mois (1-12)
    
    Returns:
        JSON: {success: bool, data: [[planning par jour]]}
    """
    try:
        sheet_name = get_month_sheet_name(year, month)
        
        wb = load_workbook()
        
        if sheet_name not in wb.sheetnames:
            return jsonify({
                'success': False,
                'error': f"Feuille '{sheet_name}' introuvable"
            }), 404
        
        sheet = wb[sheet_name]
        header_row = find_header_row(sheet, 'Matricule')
        
        if not header_row:
            return jsonify({
                'success': False,
                'error': "En-tête introuvable"
            }), 404
        
        # Lire planning
        planning_data = []
        for row_idx in range(header_row + 1, sheet.max_row + 1):
            matricule = cell_to_str(sheet.cell(row_idx, 1).value)
            
            if matricule:
                agent_data = {
                    'matricule': matricule,
                    'nom': cell_to_str(sheet.cell(row_idx, 2).value) or '',
                    'prenom': cell_to_str(sheet.cell(row_idx, 3).value) or '',
                    'jours': []
                }
                
                # Lire jours (colonnes D onwards)
                for day in range(1, 32):  # Max 31 jours
                    col_idx = PLANNING_DAYS_START_COL + day - 1
                    status = normalize_status(sheet.cell(row_idx, col_idx).value)
                    agent_data['jours'].append(status)
                
                planning_data.append(agent_data)
        
        return jsonify({
            'success': True,
            'sheet': sheet_name,
            'data': planning_data
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/planning-data/<int:year>/<int:month>', methods=['PUT'])
def update_planning_data(year, month):
    """
    Met à jour le planning pour un mois donné
    
    Args:
        year: Année
        month: Mois (1-12)
    
    Body: {data: [[planning]]}
    
    Returns:
        JSON: {success: bool, message: str}
    """
    try:
        sheet_name = get_month_sheet_name(year, month)
        data = request.get_json()
        planning_data = data.get('data', [])
        
        # Créer backup
        create_backup()
        
        # Charger workbook
        wb = load_workbook(data_only=False)
        
        if sheet_name not in wb.sheetnames:
            return jsonify({
                'success': False,
                'error': f"Feuille '{sheet_name}' introuvable"
            }), 404
        
        sheet = wb[sheet_name]
        header_row = find_header_row(sheet, 'Matricule')
        
        # Mettre à jour planning
        for i, agent_data in enumerate(planning_data):
            row_idx = header_row + 1 + i
            
            # Mettre à jour jours
            for day, status in enumerate(agent_data.get('jours', []), start=1):
                col_idx = PLANNING_DAYS_START_COL + day - 1
                sheet.cell(row_idx, col_idx, status)
        
        # Sauvegarder
        save_workbook(wb)
        
        # Upload vers SharePoint
        if AUTO_SYNC_ENABLED:
            try:
                upload_to_sharepoint()
            except SharePointError:
                pass
        
        return jsonify({
            'success': True,
            'message': f'Planning {sheet_name} mis à jour'
        })
    
    except ExcelError as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==============================================================================
# ROUTES - GÉNÉRATION
# ==============================================================================

@app.route('/api/generate-week', methods=['POST'])
def generate_week():
    """
    Génère les disponibilités pour une semaine ISO
    
    Body: {week: "2026-W01", group: "all", slots: 2}
    
    Returns:
        JSON: {success: bool, disponibilites: {}}
    """
    try:
        data = request.get_json()
        week_str = data.get('week')  # Format: "2026-W01"
        group = data.get('group', 'all')
        
        if not week_str:
            return jsonify({
                'success': False,
                'error': 'Paramètre week manquant'
            }), 400
        
        # Parser semaine ISO
        year, week_num = week_str.split('-W')
        year, week_num = int(year), int(week_num)
        
        # Calculer dates (lundi à vendredi uniquement)
        # Utiliser %G-%V pour semaine ISO correcte
        first_day = datetime.strptime(f'{year}-W{week_num:02d}-1', '%G-W%V-%u')
        days = [(first_day + timedelta(days=i)) for i in range(5)]  # 0-4 = lundi à vendredi
        
        # Charger planning
        wb = load_workbook()
        
        disponibilites = {}
        
        for day in days:
            sheet_name = get_month_sheet_name(day.year, day.month)
            
            if sheet_name not in wb.sheetnames:
                continue
            
            sheet = wb[sheet_name]
            header_row = find_header_row(sheet, 'Matricule')
            
            day_disponibles = []
            
            # Créer un mapping des agents valides
            agents_valides = {agent['nom']: agent for agent in AGENTS}
            
            for row_idx in range(header_row + 1, sheet.max_row + 1):
                nom = cell_to_str(sheet.cell(row_idx, 2).value)
                
                if not nom:
                    continue
                
                # Vérifier que c'est un vrai agent (pas un en-tête d'équipe)
                if nom not in agents_valides:
                    continue
                
                prenom = cell_to_str(sheet.cell(row_idx, 3).value)
                
                # Vérifier groupe (tous les groupes pour les disponibilités)
                if group != 'all':
                    group_agents = GROUPS.get(group, [])
                    if nom not in group_agents:
                        continue
                
                # Vérifier disponibilité
                col_idx = PLANNING_DAYS_START_COL + day.day - 1
                status = normalize_status(sheet.cell(row_idx, col_idx).value)
                
                if status not in ABSENT_STATUSES:
                    # Récupérer les permis et CACES depuis la feuille config
                    config_sheet_ref = wb['config']
                    config_header_row = find_header_row(config_sheet_ref, 'Matricule')
                    
                    # Trouver la ligne de l'agent dans config
                    agent_certs = {}
                    for cfg_row_idx in range(config_header_row + 1, config_sheet_ref.max_row + 1):
                        cfg_nom = cell_to_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_NOM).value)
                        if cfg_nom == nom:
                            agent_certs = {
                                'r482_a_1': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_A_1).value) or '',
                                'r482_b1_2': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_B1_2).value) or '',
                                'r482_c2_3': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_C2_3).value) or '',
                                'r482_c1_4': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_C1_4).value) or '',
                                'r482_5': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_5).value) or '',
                                'r482_d_7': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_D_7).value) or '',
                                'r482_e_8': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_E_8).value) or '',
                                'r482_f_9': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_F_9).value) or '',
                                'r482_g_10': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CACES_R482_G_10).value) or '',
                                'tondeuse': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_TONDEUSE).value) or '',
                                'grue_r490': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_GRUE_R490).value) or '',
                                'b_nacelle': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_B_NACELLE).value) or '',
                                'chariot_r489_3': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_CHARIOT_R489_3).value) or '',
                                'tronco': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_TRONCO).value) or '',
                                'permis_be': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_PERMIS_BE).value) or '',
                                'permis_pl': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_PERMIS_PL).value) or '',
                                'permis_ce': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_PERMIS_CE).value) or '',
                                'fimo': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_FIMO).value) or '',
                                'aipr': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_AIPR).value) or '',
                                'premiers_secours': cell_to_date_str(config_sheet_ref.cell(cfg_row_idx, AGENT_COL_PREMIERS_SECOURS).value) or ''
                            }
                            break
                    
                    day_disponibles.append({
                        'nom': nom,
                        'prenom': prenom,
                        'status': status,
                        'equipe': agents_valides[nom]['equipe'],
                        'equipe_color': TEAM_COLORS.get(agents_valides[nom]['equipe'], '#E0E0E0'),
                        'certifications': agent_certs
                    })
            
            disponibilites[day.strftime('%Y-%m-%d')] = day_disponibles
        
        return jsonify({
            'success': True,
            'week': week_str,
            'group': group,
            'disponibilites': disponibilites
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/generate-teams', methods=['POST'])
def generate_teams():
    """
    Génère les équipes par compétences (seulement Voirie et Espaces Verts)
    
    Body: {week: "2026-W01", group: "voirie" ou "espaces_verts", team_size: 3}
    
    Returns:
        JSON: {success: bool, equipes: {}}
    """
    try:
        data = request.get_json()
        week_str = data.get('week')
        group = data.get('group', 'voirie')
        team_size = data.get('team_size', 3)
        
        # Vérifier que le groupe est voirie ou espaces_verts
        if group not in ['voirie', 'espaces_verts']:
            return jsonify({
                'success': False,
                'error': 'Génération d\'équipes disponible uniquement pour "voirie" et "espaces_verts"'
            }), 400
        
        # Récupérer les disponibilités pour ce groupe
        dispo_request_data = {
            'week': week_str,
            'group': group,
            'slots': 2
        }
        
        # Simuler la requête interne
        from flask import request as flask_request
        with app.test_request_context(
            '/api/generate-week',
            method='POST',
            json=dispo_request_data
        ):
            dispo_response = generate_week()
            dispo_data = dispo_response.get_json()
        
        if not dispo_data.get('success'):
            return jsonify(dispo_data), 400
        
        disponibilites = dispo_data['disponibilites']
        
        # Générer équipes (round-robin simple)
        equipes = {}
        
        for day, agents in disponibilites.items():
            day_equipes = []
            
            # Créer équipes de taille team_size
            for i in range(0, len(agents), team_size):
                equipe = agents[i:i+team_size]
                if equipe:
                    day_equipes.append(equipe)
            
            equipes[day] = day_equipes
        
        return jsonify({
            'success': True,
            'week': week_str,
            'group': group,
            'team_size': team_size,
            'equipes': equipes
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==============================================================================
# GESTIONNAIRES D'ERREURS
# ==============================================================================

@app.errorhandler(404)
def not_found(error):
    """Gestionnaire d'erreur 404"""
    return jsonify({
        'success': False,
        'error': 'Ressource introuvable'
    }), 404

@app.errorhandler(500)
def internal_error(error):
    """Gestionnaire d'erreur 500"""
    return jsonify({
        'success': False,
        'error': 'Erreur serveur interne'
    }), 500

# ==============================================================================
# ROUTES - INFORMATIONS STATUTS
# ==============================================================================

@app.route('/api/status-info')
def get_status_info():
    """
    Retourne les informations sur tous les statuts (couleurs, libellés)
    
    Returns:
        JSON: {success: bool, statuses: {code: {color, label, is_absent}}}
    """
    try:
        status_info = {}
        
        # Tous les statuts connus
        all_statuses = set(STATUS_COLORS.keys())
        
        for status_code in all_statuses:
            status_info[status_code] = {
                'color': STATUS_COLORS.get(status_code, '#CCCCCC'),
                'label': STATUS_LABELS.get(status_code, status_code),
                'is_absent': status_code in ABSENT_STATUSES
            }
        
        return jsonify({
            'success': True,
            'statuses': status_info
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ==============================================================================
# LANCEMENT APPLICATION
# ==============================================================================

if __name__ == '__main__':
    print("\n>> Lancement ST8 Planning")
    print(f"URL: http://{FLASK_HOST}:{FLASK_PORT}")
    print(f"Fichier: {EXCEL_FILE}")
    print("\nCtrl+C pour arreter\n")
    
    # Vérifier le fichier au démarrage
    startup_sync()
    
    app.run(
        host=FLASK_HOST,
        port=FLASK_PORT,
        debug=FLASK_DEBUG
    )
